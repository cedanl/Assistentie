# /// script
# requires-python = ">=3.12"
# dependencies = [
#     "anthropic",
#     "pydantic",
#     "streamlit",
#     "streamlit_extras"
#     "pandas",
#     "requests",
#     "plotly",
#     "python-docx",
#     "Pillow",
# ]
# ///

# ─────────────────────────────────────────────────────────────────────────────
# Organization: CEDA
# Original Authors: Ed. de Feber, Edwin Lieftink, Steven Ramondt
# ─────────────────────────────────────────────────────────────────────────────

"""frontend/app.py — Streamlit frontend voor de Uitnodigingsregel app EduPlan."""

# ─────────────────────────────────────────────
# Imports
# ─────────────────────────────────────────────

import html
import streamlit as st
from streamlit_extras.bottom_container import bottom
import pandas as pd
import requests
import plotly.graph_objects as go
from concurrent.futures import ThreadPoolExecutor
from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO
from datetime import datetime
from styles import START_CSS, MAIN_CSS, TERRACOTTA, ROZE_LICHT


# ─────────────────────────────────────────────────────────────────────────────
# Paginaconfiguratie  (moet als eerste Streamlit-aanroep staan)
# ─────────────────────────────────────────────────────────────────────────────

st.set_page_config(
    layout="wide",
    initial_sidebar_state="collapsed",
    menu_items=None,
    # menu_items={
        # "Get Help": "https://github.com/cedanl/Assistentie",
        # "Report a bug": "mailto:ed.defeber@surf.nl",
        # "About": "EduPlan — CEDA 2026",
    # },
    page_icon="🧮",
    page_title="EduPlan",
)


# ─────────────────────────────────────────────
# Data & features
# ─────────────────────────────────────────────

@st.cache_data
def _load_data() -> pd.DataFrame:
    return pd.read_csv("shared/data.csv")


NON_FEATURES = {"Dropout", "Naam", "Opleiding", "Klas", "Mentor"}


# ─────────────────────────────────────────────
# Session state
# ─────────────────────────────────────────────

_defaults = {
    "page":                    "start",
    "selected_opleiding":      "Alle",
    "selected_klas":           "Alle",
    "actieve_tab":             "uitnodigingsregel",
    "toon_zoekbalk":           False,
    "top_n":                   10,
    "risicostudenten":         [],
    "filter_key":              None,
    "laatste_analyse":         None,
    "eduplan_genereren":       False,
    "geselecteerde_student":   0,
    "uploaded_df":             None,
    "gebruik_demo_data":       True,
    "upload_filename":         "",
    "toon_alle_opleidingen":   False,
    "heeft_dropout_kolom":     False,
    "training_status":         "idle",   # idle | training | done | failed
    "training_message":        "",
    "model_is_custom":         False,
}
for k, v in _defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# Gebruik geüpload bestand indien aanwezig, anders standaard data
if st.session_state.uploaded_df is not None:
    df = st.session_state.uploaded_df
else:
    df = _load_data()

features = [col for col in df.columns if col not in NON_FEATURES]
QUICK_OPLEIDINGEN = sorted(df["Opleiding"].unique().tolist())


# ─────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────

def _norm_col(s: str) -> str:
    return s.lower().replace("_", "").replace(" ", "").replace("-", "")


def _zoek_opleiding(zoekterm: str) -> str:
    return next(
        (o for o in QUICK_OPLEIDINGEN if zoekterm.lower() in o.lower()),
        zoekterm,
    )


def _pill_klik(opl: str):
    st.session_state.selected_opleiding = opl
    st.session_state.selected_klas      = "Alle"
    st.session_state.page               = "main"
    st.session_state.filter_key         = None


def _verwerk_upload(uploaded_file) -> None:
    """Laad het geüploade bestand, koppel kolommen via fuzzy-matching + LLM en vul ontbrekende aan."""
    try:
        with st.spinner("Bestand laden en kolommen koppelen…"):
            # 1. Lees bestand
            if uploaded_file.name.endswith(".xlsx"):
                new_df = pd.read_excel(uploaded_file)
            else:
                new_df = pd.read_csv(uploaded_file)

            # 2. Vereiste modelkolommen bepalen
            demo_df = _load_data()
            vereist = [c for c in demo_df.columns if c not in NON_FEATURES]

            hernoem_log: dict[str, str] = {}   # geüpload → vereist
            vul_log:    list[str]       = []

            # 2b. Detecteer uitvalkolom onder alternatieve naam (bijv. Uitval, Uitgevallen)
            if "Dropout" not in new_df.columns:
                _dropout_synoniemen = {
                    "dropout", "uitval", "uitgevallen", "uitgevalen",
                    "isuitgevallen", "uitvalindicator", "uitval_indicator",
                    "gestopt", "vroegtijdigverlaten",
                }
                _gevonden = next(
                    (c for c in new_df.columns if _norm_col(c) in _dropout_synoniemen),
                    None,
                )
                if _gevonden is None:
                    try:
                        resp = requests.post(
                            "http://localhost:8000/map_columns",
                            json={
                                "uploaded_columns": list(new_df.columns),
                                "required_columns": ["Dropout"],
                            },
                            timeout=30,
                        )
                        _llm = resp.json().get("mapping", {})
                        _kandidaat = _llm.get("Dropout")
                        if _kandidaat and _kandidaat in new_df.columns:
                            _gevonden = _kandidaat
                    except Exception:
                        pass
                if _gevonden:
                    new_df = new_df.rename(columns={_gevonden: "Dropout"})
                    hernoem_log[_gevonden] = "Dropout"

            ontbrekend = [c for c in vereist if c not in new_df.columns]

            if ontbrekend:
                # 3a. Lokale fuzzy-match (case/underscore/spatie insensitief)
                norm_upload = {_norm_col(c): c for c in new_df.columns}
                nog_ontbrekend = []
                rename_now: dict[str, str] = {}  # uploaded → required

                for req in ontbrekend:
                    kandidaat = norm_upload.get(_norm_col(req))
                    if kandidaat and kandidaat not in rename_now:
                        rename_now[kandidaat] = req
                        hernoem_log[kandidaat] = req
                    else:
                        nog_ontbrekend.append(req)

                # 3b. LLM-koppeling voor resterende ontbrekende kolommen
                if nog_ontbrekend:
                    try:
                        # Stuur alleen kolommen mee die nog niet zijn gekoppeld
                        vrije_kolommen = [c for c in new_df.columns if c not in rename_now]
                        resp = requests.post(
                            "http://localhost:8000/map_columns",
                            json={
                                "uploaded_columns": vrije_kolommen,
                                "required_columns": nog_ontbrekend,
                            },
                            timeout=30,
                        )
                        llm_mapping = resp.json().get("mapping", {})
                        for req, upl in llm_mapping.items():
                            if (
                                req in nog_ontbrekend
                                and upl in new_df.columns
                                and upl not in rename_now
                            ):
                                rename_now[upl] = req
                                hernoem_log[upl] = req
                    except Exception:
                        pass

                # 4. Kolommen hernoemen
                if rename_now:
                    new_df = new_df.rename(columns=rename_now)

                # 5. Nog steeds ontbrekend → aanvullen met mediaanwaarden uit demo-data
                nog_missen = [c for c in vereist if c not in new_df.columns]
                if nog_missen:
                    medians = demo_df[nog_missen].median()
                    for col in nog_missen:
                        new_df[col] = medians[col]
                        vul_log.append(col)

            # 6. Metadata-kolommen aanvullen indien afwezig
            if "Opleiding" not in new_df.columns:
                new_df["Opleiding"] = "Geüploade opleiding"
            if "Klas" not in new_df.columns:
                new_df["Klas"] = "Klas A"
            if "Naam" not in new_df.columns:
                if "Studentnummer" in new_df.columns:
                    new_df["Naam"] = new_df["Studentnummer"].astype(str)
                else:
                    new_df["Naam"] = [f"Lerende {i + 1}" for i in range(len(new_df))]
            if "Mentor" not in new_df.columns:
                new_df["Mentor"] = "Mentor"

            # 7. Opslaan in session state
            st.session_state.uploaded_df       = new_df
            st.session_state.upload_filename   = uploaded_file.name
            st.session_state.gebruik_demo_data = False
            st.session_state.filter_key        = None
            st.session_state.risicostudenten   = []
            st.session_state.laatste_analyse   = None

            # 8. Detecteer historische data met uitvalresultaten
            heeft_dropout = (
                "Dropout" in new_df.columns
                and new_df["Dropout"].notna().sum() >= 30
            )
            st.session_state.heeft_dropout_kolom = heeft_dropout
            if heeft_dropout:
                st.session_state.training_status = "idle"

        # 8. Feedback aan gebruiker
        regels = [f"**'{uploaded_file.name}'** geladen — {len(new_df)} lerenden."]
        if hernoem_log:
            koppelingen = ", ".join(f"`{u}` → `{r}`" for u, r in hernoem_log.items())
            regels.append(f"**Kolommen automatisch gekoppeld:** {koppelingen}")
        if vul_log:
            regels.append(
                f"**Kolommen niet gevonden, aangevuld met standaardwaarden:** "
                + ", ".join(f"`{c}`" for c in vul_log)
            )
        if hernoem_log or vul_log:
            st.info("\n\n".join(regels))
        else:
            st.success(regels[0])

        st.rerun()

    except Exception as e:
        st.error(f"Fout bij het inladen van het bestand: {e}")


def _klassen_voor(opleiding: str) -> list[str]:
    if opleiding == "Alle":
        return ["Alle"] + sorted(df["Klas"].unique().tolist())
    return ["Alle"] + sorted(df[df["Opleiding"] == opleiding]["Klas"].unique().tolist())


def _gefilterde_df():
    opl  = st.session_state.selected_opleiding
    klas = st.session_state.selected_klas
    d = df
    if opl  != "Alle": d = d[d["Opleiding"] == opl]
    if klas != "Alle": d = d[d["Klas"]      == klas]
    return d


def _build_word_doc(analyse: dict) -> BytesIO:
    doc = Document()
    doc.add_heading("EduPlan", 0)

    doc.add_heading("Studentgegevens", 1)
    info = doc.add_paragraph()
    for label, value in [
        ("Student",          analyse["naam"]),
        ("Student-ID",       str(analyse["studentnummer"])),
        ("Opleiding",        analyse["opleiding"]),
        ("Klas",             analyse["klas"]),
        ("Mentor",           analyse["mentor"]),
        ("Datum",            datetime.now().strftime("%d-%m-%Y %H:%M")),
    ]:
        info.add_run(f"{label}: ").bold = True
        info.add_run(f"{value}\n")

    doc.add_heading("Kengetallen", 2)
    kgn = doc.add_paragraph()
    kgn.add_run("Leeftijd: ").bold = True
    kgn.add_run(f"{analyse['leeftijd']} jaar\n")
    kgn.add_run("Ongeoorloofd verzuim: ").bold = True
    kgn.add_run(f"{analyse['ongeoorloofd_verzuim']:.1f} dagen\n")
    kgn.add_run("Geoorloofd verzuim: ").bold = True
    kgn.add_run(f"{analyse['geoorloofd_verzuim']:.1f} dagen\n")
    kgn.add_run("Uitvalkans: ").bold = True
    r = kgn.add_run(f"{analyse['probability']:.1%}\n")
    r.font.color.rgb = RGBColor(200, 120, 90)
    r.bold = True

    doc.add_heading("EduPlan", 1)
    doc.add_paragraph(analyse["explanation"])

    doc.add_heading("Risicofactoren (SHAP)", 1)
    doc.add_paragraph(analyse["feature_importance"])

    footer = doc.add_paragraph()
    r2 = footer.add_run("Gegenereerd door Uitnodigingsregel — CEDA")
    r2.italic = True
    r2.font.size = Pt(9)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def _run_voorspelling(dff: pd.DataFrame):
    """Doe API-calls voor alle studenten in dff; sla op in session_state."""
    opl  = st.session_state.selected_opleiding
    klas = st.session_state.selected_klas
    key  = (opl, klas)

    if st.session_state.filter_key == key:
        return  # niets veranderd

    st.session_state.filter_key      = key
    st.session_state.laatste_analyse  = None
    st.session_state.eduplan_genereren = False

    gebruik_default = st.session_state.gebruik_demo_data

    def _call(row):
        try:
            resp = requests.post(
                "http://localhost:8000/predict_dropout",
                json={
                    "student":           row[features].to_dict(),
                    "use_default_model": gebruik_default,
                },
                timeout=10,
            )
            return (row, resp.json())
        except Exception:
            return None

    with st.spinner("Risico berekenen…"):
        rows = [row for _, row in dff.iterrows()]
        with ThreadPoolExecutor(max_workers=20) as pool:
            resultaten = [r for r in pool.map(_call, rows) if r is not None]
        resultaten.sort(key=lambda x: x[1]["probability"], reverse=True)
        st.session_state.risicostudenten = resultaten
        st.session_state.top_n = min(len(resultaten), 10)


def _genereer_eduplan():
    idx    = st.session_state.geselecteerde_student
    risico = st.session_state.risicostudenten
    if not risico or idx >= len(risico):
        return
    row, result = risico[idx]
    naam = row["Naam"]

    with st.spinner(f"🕑 Bezig met genereren van het EduPlan voor {naam}…"):
        gebruik_default = st.session_state.gebruik_demo_data

        def _fetch_explain():
            try:
                return requests.post(
                    "http://localhost:8000/explain_risk",
                    json={
                        "student":           row[features].to_dict(),
                        "prediction":        result["prediction"],
                        "probability":       result["probability"],
                        "use_default_model": gebruik_default,
                    },
                    timeout=60,
                ).json()["explanation"]
            except Exception:
                return "Uitleg kon niet worden gegenereerd."

        def _fetch_fi():
            try:
                return requests.post(
                    "http://localhost:8000/feature_importance",
                    json={
                        "student":           row[features].to_dict(),
                        "use_default_model": gebruik_default,
                    },
                    timeout=10,
                ).json()["feature_importance"]
            except Exception:
                return {}

        with ThreadPoolExecutor(max_workers=2) as pool:
            f_exp = pool.submit(_fetch_explain)
            f_fi  = pool.submit(_fetch_fi)
            exp     = f_exp.result()
            fi_dict = f_fi.result()

        fi_str = ", ".join(f"{k}: {v:.2f}" for k, v in fi_dict.items()) if fi_dict else "Niet beschikbaar."

        analyse = {
            "naam":                    naam,
            "opleiding":               row.get("Opleiding", "—"),
            "klas":                    row.get("Klas", "—"),
            "mentor":                  row.get("Mentor", "—"),
            "studentnummer":           int(row["Studentnummer"]) if "Studentnummer" in row.index else "—",
            "leeftijd":                int(row["StudentAge"]) if "StudentAge" in row.index else "—",
            "ongeoorloofd_verzuim":    float(row["absence_unauthorized"]) if "absence_unauthorized" in row.index else 0.0,
            "geoorloofd_verzuim":      float(row.get("absence_authorized", 0)),
            "probability":             result["probability"],
            "explanation":             exp,
            "feature_importance":      fi_str,
            "feature_importance_dict": fi_dict,
        }
        analyse["docx"] = _build_word_doc(analyse)
        st.session_state.laatste_analyse   = analyse
        st.session_state.eduplan_genereren = False


# ─────────────────────────────────────────────
# Modeltraining
# ─────────────────────────────────────────────

def _start_training() -> None:
    upload_df = st.session_state.uploaded_df
    payload = {
        "data":           upload_df.to_dict(orient="records"),
        "dropout_column": "Dropout",
    }
    try:
        resp = requests.post("http://localhost:8000/train_model", json=payload, timeout=10)
        resultaat = resp.json().get("status")
        if resultaat in ("started", "already_running"):
            st.session_state.training_status = "training"
            st.rerun()
    except Exception as e:
        st.error(f"Kan training niet starten: {e}")


def _reset_model() -> None:
    try:
        requests.delete("http://localhost:8000/reset_model", timeout=10)
        st.session_state.training_status  = "idle"
        st.session_state.training_message = ""
        st.session_state.model_is_custom  = False
        st.session_state.risicostudenten  = []
        st.session_state.filter_key       = None
        st.rerun()
    except Exception as e:
        st.error(f"Reset mislukt: {e}")


def _show_training_panel() -> None:
    """Toont de trainings-UI op het startscherm wanneer historische data is geüpload."""
    import time

    status = st.session_state.training_status

    if status == "idle":
        st.info(
            "Je data bevat een **Dropout-kolom** met historische uitvalgegevens. "
            "Train het model op jouw eigen data voor instelling-specifieke voorspellingen."
        )
        if st.button("Train model op jouw data", type="primary", use_container_width=True):
            _start_training()

    elif status == "training":
        # Bewaar starttijd in session state — overleeft Streamlit-reruns
        if "training_start_time" not in st.session_state:
            st.session_state.training_start_time = time.time()

        verstreken = time.time() - st.session_state.training_start_time
        minuten, seconden = divmod(int(verstreken), 60)
        tijdtekst = f"{minuten}m {seconden}s" if minuten else f"{seconden}s"

        stappen = [
            (0,  "Gegevens laden en features valideren…"),
            (5,  "Hyperparameterraster opbouwen (24 combinaties × 5-fold CV)…"),
            (15, "Modellen fitten — dit duurt het langst…"),
            (45, "Laatste fits afronden en beste model selecteren…"),
        ]
        fase = next(
            (msg for drempel, msg in reversed(stappen) if verstreken >= drempel),
            stappen[0][1],
        )

        with st.status("Model wordt getraind…", expanded=True):
            st.write("GridSearchCV doorzoekt hyperparameters. Dit duurt doorgaans 15–60 seconden.")
            st.markdown(f"⏱ **Verstreken tijd:** {tijdtekst}")
            st.caption(f"Fase: {fase}")

        # Eén statuscheck per render — geen blokkerende loop
        try:
            data = requests.get("http://localhost:8000/train_status", timeout=5).json()
        except Exception:
            data = {"status": "training"}

        if data["status"] == "done":
            totaal = int(time.time() - st.session_state.pop("training_start_time", time.time()))
            st.session_state.training_status  = "done"
            st.session_state.training_message = f"{data['message']} (getraind in {totaal}s)"
            st.session_state.model_is_custom  = True
            st.session_state.risicostudenten  = []
            st.session_state.filter_key       = None
            st.rerun()
        elif data["status"] == "failed":
            st.session_state.pop("training_start_time", None)
            st.session_state.training_status  = "failed"
            st.session_state.training_message = data["message"]
            st.rerun()
        else:
            # Nog bezig — wacht kort en herlaad
            time.sleep(2)
            st.rerun()

    elif status == "done":
        bericht = st.session_state.get("training_message", "")
        st.success(
            f"Instellingsmodel actief. {bericht} "
            "Voorspellingen worden gedaan met jouw eigen getrainde model."
        )
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Terugzetten naar standaardmodel", use_container_width=True):
                _reset_model()
        with col2:
            st.caption("Het model blijft actief na herstart van de backend.")

    elif status == "failed":
        bericht = st.session_state.get("training_message", "")
        st.error(f"Training mislukt: {bericht}")
        if st.button("Opnieuw proberen", use_container_width=True):
            st.session_state.training_status = "idle"
            st.rerun()


# ─────────────────────────────────────────────
# Startscherm
# ─────────────────────────────────────────────

def show_start_screen():
    st.markdown(START_CSS, unsafe_allow_html=True)
    st.markdown("<div style='height:60px'></div>", unsafe_allow_html=True)

    # ── Titel ──
    st.markdown(
        """
        <div style="text-align:center; margin-bottom:32px; font-family:'General Sans',sans-serif;">
            <h1 style="font-size:3.2rem; font-weight:600; line-height:1.15; margin-bottom:4px; padding:0;">
                Welkom bij de<br>Uitnodigingsregel
            </h1>
            <p style="font-size:1.3rem; font-weight:500; color:#333; margin-top:0; padding:0;">
                op tijd de juiste lerenden uitnodigen
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    _, col_m, _ = st.columns([1, 4, 1])
    with col_m:
        # ── Beschrijvingstekst ──
        st.markdown(
            """<p style="text-align:center; font-size:1rem; color:#333; margin-bottom:20px;
                         font-family:'General Sans',sans-serif; line-height:1.6;">
                Voeg je eigen dataset met lerenden toe, om te zien of er nú lerenden zijn
                die mogelijk risico lopen om uit te vallen. We brengen ze voor jou in beeld.
            </p>""",
            unsafe_allow_html=True,
        )

        # ── Bestand uploaden ──
        uploaded_file = st.file_uploader(
            "Voeg je databestand toe",
            type=["csv", "xlsx"],
            label_visibility="collapsed",
        )

        # Verwerk alleen nieuw geüpload bestand (voorkomt herhaald verwerken na rerun)
        if uploaded_file is not None and uploaded_file.name != st.session_state.upload_filename:
            _verwerk_upload(uploaded_file)

        # ── Modeltraining (alleen bij geüploade historische data met Dropout-kolom) ──
        if (
            st.session_state.heeft_dropout_kolom
            and not st.session_state.gebruik_demo_data
        ):
            _show_training_panel()
            st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

        # ── Demo-data checkbox ──
        demo_nieuw = st.checkbox(
            "Of gebruik de synthetische demo-data",
            value=st.session_state.gebruik_demo_data,
        )
        if demo_nieuw != st.session_state.gebruik_demo_data:
            st.session_state.gebruik_demo_data = demo_nieuw
            if demo_nieuw:
                st.session_state.uploaded_df     = None
                st.session_state.upload_filename = ""
                st.session_state.filter_key      = None
                st.session_state.risicostudenten = []
                st.session_state.laatste_analyse = None

        st.markdown("<div style='height:16px'></div>", unsafe_allow_html=True)

        # ── START-knop ──
        data_beschikbaar = st.session_state.gebruik_demo_data or st.session_state.uploaded_df is not None
        if st.button(
            "START DE UITNODIGINGSREGEL",
            type="primary",
            use_container_width=True,
            disabled=not data_beschikbaar,
        ):
            st.session_state.page               = "main"
            st.session_state.selected_opleiding = "Alle"
            st.session_state.selected_klas      = "Alle"
            st.rerun()

        if not data_beschikbaar:
            st.caption("Upload een databestand of selecteer de synthetische demo-data om te starten.")

    st.markdown("<div style='height:20px'></div>", unsafe_allow_html=True)

    # ── Opleidingen-pills ──
    _, col_pills, _ = st.columns([0.5, 6, 0.5])
    with col_pills:
        ZICHTBAAR  = 4
        eerste_rij = QUICK_OPLEIDINGEN[:ZICHTBAAR]
        rest       = QUICK_OPLEIDINGEN[ZICHTBAAR:]

        pill_cols = st.columns(ZICHTBAAR + 1)
        for i, opl in enumerate(eerste_rij):
            with pill_cols[i]:
                if st.button(opl, key=f"pill_{opl}", use_container_width=True):
                    _pill_klik(opl)
                    st.rerun()
        with pill_cols[ZICHTBAAR]:
            if not st.session_state.toon_alle_opleidingen:
                if st.button("Meer ↓", key="pill_meer", use_container_width=True):
                    st.session_state.toon_alle_opleidingen = True
                    st.rerun()
            else:
                if st.button("Minder ↑", key="pill_minder", use_container_width=True):
                    st.session_state.toon_alle_opleidingen = False
                    st.rerun()

        if st.session_state.toon_alle_opleidingen and rest:
            extra_cols = st.columns(min(len(rest), 5))
            for i, opl in enumerate(rest):
                with extra_cols[i % len(extra_cols)]:
                    if st.button(opl, key=f"pill_extra_{opl}", use_container_width=True):
                        _pill_klik(opl)
                        st.rerun()


# ─────────────────────────────────────────────
# Hoofdscherm — header
# ─────────────────────────────────────────────

def _render_header():
    tab = st.session_state.actieve_tab

    col_ceda, col_terug, col_ur, col_ep = st.columns([3, 1.2, 2.2, 1.2])

    with col_ceda:
        st.markdown(
            "<p style='font-weight:700;font-size:1.5rem;font-family:\"General Sans\",sans-serif;"
            "margin:0;padding:6px 0;'>CEDA</p>",
            unsafe_allow_html=True,
        )
    with col_terug:
        st.markdown("<div class='nav-terug'>", unsafe_allow_html=True)
        if st.button("← TERUG", key="nav_terug", use_container_width=True):
            st.session_state.page = "start"
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)
    with col_ur:
        klasse_ur = "nav-actief" if tab == "uitnodigingsregel" else "nav-inactief"
        st.markdown(f"<div class='{klasse_ur}'>", unsafe_allow_html=True)
        if st.button("UITNODIGINGSREGEL", key="nav_ur", use_container_width=True):
            st.session_state.actieve_tab = "uitnodigingsregel"
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)
    with col_ep:
        klasse_ep = "nav-actief" if tab == "eduplan" else "nav-inactief"
        st.markdown(f"<div class='{klasse_ep}'>", unsafe_allow_html=True)
        if st.button("EDUPLAN", key="nav_ep", use_container_width=True):
            st.session_state.actieve_tab = "eduplan"
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)


# ─────────────────────────────────────────────
# Hoofdscherm — kaart-header (opleiding + klas + potlood)
# ─────────────────────────────────────────────

def _render_card_header():
    opl = st.session_state.selected_opleiding

    if st.session_state.toon_zoekbalk:
        st.markdown("<div class='card-zoek'>", unsafe_allow_html=True)
        col_s, col_b = st.columns([5, 1])
        with col_s:
            zoek = st.text_input(
                "zoek",
                placeholder="🔍  Zoek een andere opleiding (b.v. Economie, Techniek)...",
                label_visibility="collapsed",
                key="card_zoek_input",
            )
        with col_b:
            if st.button("ZOEK", type="primary", use_container_width=True, key="card_zoek_btn"):
                if zoek.strip():
                    st.session_state.selected_opleiding = _zoek_opleiding(zoek.strip())
                st.session_state.selected_klas       = "Alle"
                st.session_state.toon_zoekbalk        = False
                st.session_state.filter_key           = None
                st.session_state.risicostudenten      = []
                st.session_state.laatste_analyse      = None
                st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    else:
        klassen = _klassen_voor(opl)

        if st.session_state.selected_klas not in klassen:
            st.session_state.selected_klas = "Alle"

        col_t, col_k, col_sp, col_p = st.columns([3, 2.5, 2.5, 0.5])

        with col_t:
            badge = (
                " <span style='font-size:0.7rem; background:#e8f5e9; color:#2e7d32;"
                " border-radius:4px; padding:2px 6px; vertical-align:middle;"
                " font-weight:500;'>instellingsmodel</span>"
                if st.session_state.get("model_is_custom") else ""
            )
            st.markdown(
                f"<h3 style='font-weight:500; margin:0; padding:6px 0;"
                f"font-family:\"General Sans\",sans-serif;'>{opl}{badge}</h3>",
                unsafe_allow_html=True,
            )

        with col_k:
            klas_idx = klassen.index(st.session_state.selected_klas) \
                if st.session_state.selected_klas in klassen else 0
            gekozen_klas = st.selectbox(
                "klas",
                klassen,
                index=klas_idx,
                label_visibility="collapsed",
                format_func=lambda x: f"KLAS:  {x}" if x != "Alle" else "KLAS:  Alle",
                key="klas_dropdown",
            )
            if gekozen_klas != st.session_state.selected_klas:
                st.session_state.selected_klas   = gekozen_klas
                st.session_state.filter_key      = None
                st.session_state.laatste_analyse = None
                st.rerun()

        with col_p:
            st.markdown("<div class='potlood-btn'>", unsafe_allow_html=True)
            if st.button("🖊️ ", key="potlood"):
                st.session_state.toon_zoekbalk = True
                st.rerun()
            st.markdown("</div>", unsafe_allow_html=True)


# ─────────────────────────────────────────────
# Hoofdscherm — banner  "Toon mij X lerenden…"
# ─────────────────────────────────────────────

def _render_banner():
    risico = st.session_state.risicostudenten

    label = f"<u><b>{st.session_state.top_n}</b></u>" if risico else "···"

    st.markdown(
        f"""<div style="background:{TERRACOTTA}; border-radius:12px; padding:14px 24px;
                        color:white; font-size:16px; font-weight:500;
                        text-align:center; margin:16px 0;
                        font-family:'General Sans',sans-serif;">
            Toon mij {label} lerenden met het hoogste risico om uit te vallen.
        </div>""",
        unsafe_allow_html=True,
    )

    if risico:
        max_n = len(risico)
        nieuw_n = st.slider(
            "Aantal te tonen",
            min_value=1,
            max_value=max_n,
            value=min(st.session_state.top_n, max_n),
            label_visibility="collapsed",
            key="top_n_slider",
        )
        if nieuw_n != st.session_state.top_n:
            st.session_state.top_n = nieuw_n
            st.rerun()


# ─────────────────────────────────────────────
# Hoofdscherm — barchart
# ─────────────────────────────────────────────

def _render_barchart():
    risico = st.session_state.risicostudenten

    if not risico:
        st.markdown(
            f"""<div style="background:{ROZE_LICHT}; border-radius:14px;
                            height:260px; display:flex; align-items:center;
                            justify-content:center; color:#bbb; font-size:1.4rem;
                            margin:8px 0; letter-spacing:0.1em;">
                · · ·
            </div>""",
            unsafe_allow_html=True,
        )
        return

    top_n = st.session_state.top_n
    top   = risico[:top_n]
    n     = len(top)

    namen  = [row["Naam"]             for row, _      in reversed(top)]
    kansen = [result["probability"]   for _,   result in reversed(top)]

    def terracotta(i, total):
        t = i / max(total - 1, 1)
        r = int(0xa0 + (0xdf - 0xa0) * t)
        g = int(0x55 + (0x9a - 0x55) * t)
        b = int(0x35 + (0x75 - 0x35) * t)
        return f"rgb({r},{g},{b})"

    kleuren = [terracotta(i, n) for i in range(n)]

    fig = go.Figure(go.Bar(
        x=kansen,
        y=namen,
        orientation="h",
        marker_color=kleuren,
        text=[f"{k:.0%}" for k in kansen],
        textposition="outside",
        textfont=dict(size=14, color="#aaa"),
    ))
    x_max = max(kansen) * 1.35 if kansen else 1.0
    fig.update_layout(
        xaxis=dict(range=[0, x_max], showgrid=False, showticklabels=False, zeroline=False),
        yaxis=dict(showgrid=False, tickfont=dict(size=14, color="#1a1a1a")),
        plot_bgcolor="white",
        paper_bgcolor="white",
        margin=dict(l=20, r=60, t=12, b=12),
        height=max(180, 60 + 56 * n),
        showlegend=False,
    )
    st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})


# ─────────────────────────────────────────────
# Hoofdscherm — EduPlan-sectie
# ─────────────────────────────────────────────

def _render_eduplan_sectie():
    risico = st.session_state.risicostudenten
    top_n  = st.session_state.top_n

    if not risico:
        st.info("Laad eerst studenten via het UITNODIGINGSREGEL-tabblad.")
        return

    top = risico[:top_n]

    st.markdown(
        "<p style='font-size:14px; color:#000; margin:8px 0 4px 0;"
        "font-family:\"General Sans\",sans-serif;'>"
        "<b>Selecteer een lerende voor de uitleg van diens uitvalrisico</b></p>",
        unsafe_allow_html=True,
    )

    opties = [
        f"{row['Naam'].upper()} — UITVALKANS: {result['probability']:.0%}"
        for row, result in top
    ]

    col_sel, col_btn = st.columns([4, 2])
    with col_sel:
        st.markdown("<div class='student-sel'>", unsafe_allow_html=True)
        idx = st.selectbox(
            "student",
            options=range(len(opties)),
            format_func=lambda i: opties[i],
            label_visibility="collapsed",
            key="student_selector",
        )
        st.markdown("</div>", unsafe_allow_html=True)

    with col_btn:
        if st.button("TOON EDUPLAN", type="primary", use_container_width=True):
            st.session_state.geselecteerde_student = idx
            st.session_state.eduplan_genereren     = True
            st.session_state.laatste_analyse       = None
            st.session_state.actieve_tab           = "eduplan"
            st.rerun()

    if st.session_state.eduplan_genereren:
        naam = html.escape(top[st.session_state.geselecteerde_student][0]["Naam"])
        st.markdown(
            f"""<div style="background:{ROZE_LICHT}; border-radius:14px;
                            padding:28px; text-align:center; margin-top:16px;
                            font-family:'General Sans',sans-serif; color:#555;">
                <span style="font-size:1.4rem;">↻</span>&nbsp;&nbsp;
                Het EduPlan voor <b>{naam}</b> wordt gemaakt
            </div>""",
            unsafe_allow_html=True,
        )
        _genereer_eduplan()
        st.rerun()

    if st.session_state.laatste_analyse:
        _render_eduplan_content()


def _render_eduplan_content():
    analyse = st.session_state.laatste_analyse
    naam    = html.escape(analyse["naam"])

    with st.container(border=True):
        st.markdown(
            f"""<div style="display:flex; align-items:center; gap:14px;
                            font-family:'General Sans',sans-serif;">
                <div style="background:{ROZE_LICHT}; color:#1a1a1a; border-radius:8px;
                            width:36px; height:36px; display:flex; align-items:center;
                            justify-content:center; font-size:16px; font-weight:700;
                            flex-shrink:0;">🚦</div>
                <span style="font-size:1.25rem; font-weight:600;">EduPlan | {naam}</span>
            </div>""",
            unsafe_allow_html=True,
        )

    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

    st.markdown(
        f"<div style='font-family:\"General Sans\",sans-serif; font-size:15px; "
        f"line-height:1.85; background:{ROZE_LICHT}; border-radius:16px; "
        f"padding:28px 32px;'>"
        f"{analyse['explanation']}"
        f"</div>",
        unsafe_allow_html=True,
    )

    st.markdown("<div style='height:16px'></div>", unsafe_allow_html=True)

    _, col_p, col_d = st.columns([4, 1.5, 1.5])

    with col_p:
        st.markdown(
            """<button onclick="(window.parent||window).print()"
                style="background-color:#1a1a1a;color:white;border-radius:50px;
                       border:none;font-size:12px;font-weight:700;letter-spacing:0.07em;
                       padding:8px 24px;box-shadow:none;cursor:pointer;width:100%;
                       white-space:nowrap;font-family:'General Sans',sans-serif;">
                PRINT
            </button>""",
            unsafe_allow_html=True,
        )

    with col_d:
        st.download_button(
            label="DOWNLOAD",
            data=analyse["docx"],
            file_name=f"EduPlan_{naam.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="download_btn",
        )


# ─────────────────────────────────────────────
# Hoofdscherm — footer
# ─────────────────────────────────────────────

def _render_footer():
    st.markdown("<div style='height:32px'></div>", unsafe_allow_html=True)
    st.markdown(
        """<hr style="border:none; border-top:2px solid rgba(0,0,0,0.32); margin:0 0 20px 0;">
        <p style="text-align:center; font-size:0.75rem; font-weight:500; color:gray; font-family:'General Sans',sans-serif;">
            &#169; &#9432; 2026 Op deze analytics tool is de Creative Commons ShareAlike
            Naamsvermelding 4.0-licentie van toepassing. Maak bij gebruik van dit werk
            vermelding van de volgende referentie: AI en data waarde(n)vol inzetten: CEDA.
            Uitnodigingsregel – EduPlan. Utrecht: Npuls
        </p>""",
        unsafe_allow_html=True,
    )
# ─────────────────────────────────────────────
# Hoofdscherm — samenstellen
# ─────────────────────────────────────────────

def show_main_screen():
    st.markdown(MAIN_CSS, unsafe_allow_html=True)

    _render_header()

    dff = _gefilterde_df()

    if len(dff) > 0:
        _run_voorspelling(dff)

    with st.container(border=True):
        _render_card_header()

        tab = st.session_state.actieve_tab

        _render_banner()

        if tab == "uitnodigingsregel":
            _render_barchart()
        else:
            _render_eduplan_sectie()

    # _render_footer()


# ─────────────────────────────────────────────
# Router
# ─────────────────────────────────────────────

if st.session_state.page == "start":
    show_start_screen()
else:
    show_main_screen()

with bottom():
    _col_txt, _col_btn = st.columns([5, 1])
    with _col_txt:
        st.markdown(
            """<p style="text-align:center; font-size:0.5rem; font-weight:500; color:#1a1a1a;
                    font-family:'General Sans',sans-serif; margin:1px 0;"><br>
                    <img src="https://mirrors.creativecommons.org/presskit/icons/cc.svg" alt="" style="max-width: 2em;max-height:3em;margin-left: .2em;"><img src="https://mirrors.creativecommons.org/presskit/icons/by.svg" alt="" style="max-width: 2em;max-height:3em;margin-left: .2em;"> Op deze analytics tool is de Creative Commons ShareAlike
            Naamsvermelding 4.0-licentie van toepassing. <br>Maak bij gebruik van dit werk
            vermelding van de volgende referentie: AI en data waarde(n)vol inzetten: CEDA.
            2026 Uitnodigingsregel – EduPlan. Utrecht: Npuls
        </p>""",
            unsafe_allow_html=True,
        )
    with _col_btn:
        if st.session_state.page != "start":
            try:
                with open("docs/model_analysis.html", "rb") as _f:
                    st.download_button(
                        label="📊 Download model evaluatie",
                        data=_f.read(),
                        file_name="model_analysis.html",
                        mime="text/html",
                        use_container_width=True,
                    )
            except FileNotFoundError:
                pass
