# /// script
# requires-python = ">=3.12"
# dependencies = [
#     "anthropic",
#     "pydantic",
#     "streamlit",
#     "streamlit_extras"
#     "pandas",
#     "requests",
#     "plotly",
#     "python-docx",
#     "Pillow",
# ]
# ///

# ─────────────────────────────────────────────────────────────────────────────
# Organization: CEDA
# Original Authors: Ed. de Feber, Edwin Lieftink, Steven Ramondt
# ─────────────────────────────────────────────────────────────────────────────

"""frontend/app.py — Streamlit frontend voor de Uitnodigingsregel app EduPlan."""

# ─────────────────────────────────────────────
# Imports
# ─────────────────────────────────────────────

import html
import json
import time
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime
from io import BytesIO

import pandas as pd
import plotly.graph_objects as go
import requests
import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor
from streamlit_extras.bottom_container import bottom
from styles import MAIN_CSS, ROZE_LICHT, START_CSS, TERRACOTTA

# ─────────────────────────────────────────────────────────────────────────────
# Paginaconfiguratie  (moet als eerste Streamlit-aanroep staan)
# ─────────────────────────────────────────────────────────────────────────────

st.set_page_config(
    layout="wide",
    initial_sidebar_state="collapsed",
    menu_items=None,
    # menu_items={
    # "Get Help": "https://github.com/cedanl/Assistentie",
    # "Report a bug": "mailto:ed.defeber@surf.nl",
    # "About": "EduPlan — CEDA 2026",
    # },
    page_icon="🧮",
    page_title="EduPlan",
)


# ─────────────────────────────────────────────
# Data & features
# ─────────────────────────────────────────────


@st.cache_data
def _load_data() -> pd.DataFrame:
    return pd.read_csv("shared/data.csv")


NON_FEATURES = {"Dropout", "Naam", "Opleiding", "Klas", "Mentor"}


# ─────────────────────────────────────────────
# Session state
# ─────────────────────────────────────────────

_defaults = {
    "page": "start",
    "selected_opleiding": "Alle",
    "selected_klas": "Alle",
    "actieve_tab": "uitnodigingsregel",
    "toon_zoekbalk": False,
    "top_n": 10,
    "risicostudenten": [],
    "filter_key": None,
    "laatste_analyse": None,
    "eduplan_genereren": False,
    "geselecteerde_student": 0,
    "trainings_df": None,  # historische data met Dropout-kolom (voor training)
    "trainings_filename": "",
    "predictie_df": None,  # huidig cohort zonder Dropout (voor ranking)
    "predictie_filename": "",
    "gebruik_demo_data": True,
    "toon_alle_opleidingen": False,
    "heeft_dropout_kolom": False,
    "training_status": "idle",  # idle | training | done | failed
    "training_message": "",
    "model_is_custom": False,
    "vul_log": [],
}
for k, v in _defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# Gebruik huidig cohort indien aanwezig, anders standaard demo-data
if st.session_state.predictie_df is not None:
    df = st.session_state.predictie_df
else:
    df = _load_data()

QUICK_OPLEIDINGEN = sorted(df["Opleiding"].unique().tolist())


# ─────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────


def _norm_col(s: str) -> str:
    return s.lower().replace("_", "").replace(" ", "").replace("-", "")


def _zoek_opleiding(zoekterm: str) -> str:
    return next(
        (o for o in QUICK_OPLEIDINGEN if zoekterm.lower() in o.lower()),
        zoekterm,
    )


def _pill_klik(opl: str) -> None:
    st.session_state.selected_opleiding = opl
    st.session_state.selected_klas = "Alle"
    st.session_state.page = "main"
    st.session_state.filter_key = None


def _lees_bestand(uploaded_file) -> pd.DataFrame:
    """Lees CSV of Excel naar DataFrame."""
    if uploaded_file.name.endswith(".xlsx"):
        return pd.read_excel(uploaded_file)
    return pd.read_csv(uploaded_file, sep=None, engine="python")


def _match_kolommen(
    new_df: pd.DataFrame,
) -> tuple[pd.DataFrame, dict[str, str], list[str]]:
    """Match kolommen van geüpload bestand naar model-features via fuzzy-match + LLM.

    Retourneert (aangepaste df, hernoem_log, vul_log).
    """
    demo_df = _load_data()
    vereist = [c for c in demo_df.columns if c not in NON_FEATURES]
    hernoem_log: dict[str, str] = {}
    vul_log: list[str] = []

    ontbrekend = [c for c in vereist if c not in new_df.columns]
    if ontbrekend:
        # Lokale fuzzy-match
        norm_upload = {_norm_col(c): c for c in new_df.columns}
        nog_ontbrekend = []
        rename_now: dict[str, str] = {}

        for req in ontbrekend:
            kandidaat = norm_upload.get(_norm_col(req))
            if kandidaat and kandidaat not in rename_now:
                rename_now[kandidaat] = req
                hernoem_log[kandidaat] = req
            else:
                nog_ontbrekend.append(req)

        # LLM-koppeling voor resterende kolommen
        if nog_ontbrekend:
            try:
                vrije_kolommen = [c for c in new_df.columns if c not in rename_now]
                resp = requests.post(
                    "http://localhost:8000/map_columns",
                    json={"uploaded_columns": vrije_kolommen, "required_columns": nog_ontbrekend},
                    timeout=30,
                )
                for req, upl in resp.json().get("mapping", {}).items():
                    if req in nog_ontbrekend and upl in new_df.columns and upl not in rename_now:
                        rename_now[upl] = req
                        hernoem_log[upl] = req
            except Exception:
                pass

        if rename_now:
            new_df = new_df.rename(columns=rename_now)

        # Nog steeds ontbrekend → aanvullen met mediaanwaarden
        nog_missen = [c for c in vereist if c not in new_df.columns]
        if nog_missen:
            medians = demo_df[nog_missen].median()
            for col in nog_missen:
                new_df[col] = medians[col]
                vul_log.append(col)

    # Metadata-kolommen aanvullen indien afwezig
    if "Opleiding" not in new_df.columns:
        new_df["Opleiding"] = "Geüploade opleiding"
    if "Klas" not in new_df.columns:
        new_df["Klas"] = "Klas A"
    if "Naam" not in new_df.columns:
        new_df["Naam"] = (
            new_df["Studentnummer"].apply(lambda x: f"Student {x}")
            if "Studentnummer" in new_df.columns
            else [f"Lerende {i + 1}" for i in range(len(new_df))]
        )
    if "Mentor" not in new_df.columns:
        new_df["Mentor"] = "Mentor"

    return new_df, hernoem_log, vul_log


def _feedback_upload(filename: str, n: int, hernoem_log: dict, vul_log: list) -> None:
    """Toon feedback na succesvolle upload."""
    regels = [f"**'{filename}'** geladen — {n} lerenden."]
    if hernoem_log:
        koppelingen = ", ".join(f"`{u}` → `{r}`" for u, r in hernoem_log.items())
        regels.append(f"**Kolommen automatisch gekoppeld:** {koppelingen}")
    if vul_log:
        regels.append(
            "**Kolommen niet gevonden, aangevuld met standaardwaarden:** " + ", ".join(f"`{c}`" for c in vul_log)
        )
    if hernoem_log or vul_log:
        st.info("\n\n".join(regels))
    else:
        st.success(regels[0])


def _verwerk_trainingsdata(uploaded_file) -> None:
    """Laad historische data (met Dropout-kolom) voor modeltraining."""
    try:
        with st.spinner("Trainingsdata laden en kolommen koppelen…"):
            new_df = _lees_bestand(uploaded_file)

            # Detecteer uitvalkolom onder alternatieve naam
            if "Dropout" not in new_df.columns:
                _synoniemen = {
                    "dropout",
                    "uitval",
                    "uitgevallen",
                    "uitgevalen",
                    "isuitgevallen",
                    "uitvalindicator",
                    "uitval_indicator",
                    "gestopt",
                    "vroegtijdigverlaten",
                }
                _gevonden = next((c for c in new_df.columns if _norm_col(c) in _synoniemen), None)
                if _gevonden is None:
                    try:
                        resp = requests.post(
                            "http://localhost:8000/map_columns",
                            json={"uploaded_columns": list(new_df.columns), "required_columns": ["Dropout"]},
                            timeout=30,
                        )
                        _kandidaat = resp.json().get("mapping", {}).get("Dropout")
                        if _kandidaat and _kandidaat in new_df.columns:
                            _gevonden = _kandidaat
                    except Exception:
                        pass
                if _gevonden:
                    new_df = new_df.rename(columns={_gevonden: "Dropout"})

            new_df, hernoem_log, vul_log = _match_kolommen(new_df)

            heeft_dropout = "Dropout" in new_df.columns and new_df["Dropout"].notna().sum() >= 30
            st.session_state.trainings_df = new_df
            st.session_state.trainings_filename = uploaded_file.name
            st.session_state.heeft_dropout_kolom = heeft_dropout
            if heeft_dropout:
                st.session_state.training_status = "idle"

        _feedback_upload(uploaded_file.name, len(new_df), hernoem_log, vul_log)
        st.rerun()
    except Exception as e:
        st.error(f"Fout bij het inladen van de trainingsdata: {e}")


def _verwerk_predictiedata(uploaded_file) -> None:
    """Laad huidig cohort (zonder Dropout-kolom) voor ranking."""
    try:
        with st.spinner("Huidig cohort laden en kolommen koppelen…"):
            new_df = _lees_bestand(uploaded_file)
            new_df, hernoem_log, vul_log = _match_kolommen(new_df)

            st.session_state.predictie_df = new_df
            st.session_state.predictie_filename = uploaded_file.name
            st.session_state.gebruik_demo_data = False
            st.session_state.vul_log = vul_log
            st.session_state.filter_key = None
            st.session_state.risicostudenten = []
            st.session_state.laatste_analyse = None

        _feedback_upload(uploaded_file.name, len(new_df), hernoem_log, vul_log)
        st.rerun()
    except Exception as e:
        st.error(f"Fout bij het inladen van het huidig cohort: {e}")


def _klassen_voor(opleiding: str) -> list[str]:
    if opleiding == "Alle":
        return ["Alle"] + sorted(df["Klas"].unique().tolist())
    return ["Alle"] + sorted(df[df["Opleiding"] == opleiding]["Klas"].unique().tolist())


def _gefilterde_df() -> pd.DataFrame:
    opl = st.session_state.selected_opleiding
    klas = st.session_state.selected_klas
    d = df
    if opl != "Alle":
        d = d[d["Opleiding"] == opl]
    if klas != "Alle":
        d = d[d["Klas"] == klas]
    return d


def _build_word_doc(analyse: dict) -> BytesIO:
    doc = Document()
    doc.add_heading("EduPlan", 0)

    doc.add_heading("Studentgegevens", 1)
    info = doc.add_paragraph()
    for label, value in [
        ("Student", analyse["naam"]),
        ("Student-ID", str(analyse["studentnummer"])),
        ("Opleiding", analyse["opleiding"]),
        ("Klas", analyse["klas"]),
        ("Mentor", analyse["mentor"]),
        ("Datum", datetime.now().strftime("%d-%m-%Y %H:%M")),
    ]:
        info.add_run(f"{label}: ").bold = True
        info.add_run(f"{value}\n")

    doc.add_heading("Kengetallen", 2)
    kgn = doc.add_paragraph()
    kgn.add_run("Leeftijd: ").bold = True
    _leeftijd = analyse["leeftijd"]
    kgn.add_run(f"{_leeftijd} jaar\n" if _leeftijd != "niet beschikbaar" else "niet beschikbaar\n")
    kgn.add_run("Ongeoorloofd verzuim: ").bold = True
    _ongeoorl = analyse["ongeoorloofd_verzuim"]
    kgn.add_run(f"{_ongeoorl:.1f} dagen\n" if isinstance(_ongeoorl, float) else f"{_ongeoorl}\n")
    kgn.add_run("Geoorloofd verzuim: ").bold = True
    _geoorl = analyse["geoorloofd_verzuim"]
    kgn.add_run(f"{_geoorl:.1f} dagen\n" if isinstance(_geoorl, float) else f"{_geoorl}\n")
    kgn.add_run("Uitvalkans: ").bold = True
    r = kgn.add_run(f"{analyse['probability']:.1%}\n")
    r.font.color.rgb = RGBColor(200, 120, 90)
    r.bold = True

    doc.add_heading("EduPlan", 1)
    doc.add_paragraph(analyse["explanation"])

    doc.add_heading("Risicofactoren (SHAP)", 1)
    doc.add_paragraph(analyse["feature_importance"])

    footer = doc.add_paragraph()
    r2 = footer.add_run("Gegenereerd door Uitnodigingsregel — CEDA")
    r2.italic = True
    r2.font.size = Pt(9)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def _run_voorspelling(dff: pd.DataFrame):
    """Doe API-call voor alle studenten in dff; sla op in session_state."""
    opl = st.session_state.selected_opleiding
    klas = st.session_state.selected_klas
    key = (opl, klas)

    if st.session_state.filter_key == key:
        return  # niets veranderd

    st.session_state.filter_key = key
    st.session_state.laatste_analyse = None
    st.session_state.eduplan_genereren = False

    gebruik_default = st.session_state.gebruik_demo_data

    with st.spinner("Risico berekenen…"):
        resp = requests.post(
            "http://localhost:8000/rank_students",
            json={
                "students": dff.to_dict(orient="records"),
                "use_default_model": gebruik_default,
            },
            timeout=30,
        )
        ranked = resp.json()
        resultaten = [(pd.Series(s), {"probability": s["probability"], "prediction": s["prediction"]}) for s in ranked]
        st.session_state.risicostudenten = resultaten
        st.session_state.top_n = min(len(resultaten), 10)


def _genereer_eduplan():
    idx = st.session_state.geselecteerde_student
    risico = st.session_state.risicostudenten
    if not risico or idx >= len(risico):
        return
    row, result = risico[idx]
    naam = row["Naam"]

    gebruik_default = st.session_state.gebruik_demo_data
    vul_log = st.session_state.get("vul_log", [])
    vul_set = set(vul_log)

    def _error_html(title: str, detail: str) -> str:
        return (
            "<div style='border-left:4px solid #c0392b; background:#fdf3f1; "
            "padding:14px 18px; border-radius:6px;'>"
            f"<b style='color:#c0392b;'>⚠️ {html.escape(title)}</b><br>"
            f"<span style='color:#444;'>{html.escape(detail)}</span>"
            "</div>"
        )

    def _fetch_fi():
        # Stille fallback naar {}: feature importance voedt een ondersteunende
        # bar chart die zelf gracefully omgaat met een lege dict. Een fout hier
        # mag het hoofd-EduPlan niet overschaduwen.
        try:
            resp = requests.post(
                "http://localhost:8000/feature_importance",
                json={
                    "student": row.to_dict(),
                    "use_default_model": gebruik_default,
                },
                timeout=10,
            )
            if resp.status_code != 200:
                return {}
            return resp.json().get("feature_importance", {})
        except (requests.RequestException, ValueError):
            return {}

    # Feature importance loopt concurrent in een thread (geen st.* erin, dus veilig
    # buiten de main-thread); de streaming-render hieronder MOET op de main thread.
    pool = ThreadPoolExecutor(max_workers=1)
    f_fi = pool.submit(_fetch_fi)

    # ── Stream: Sectie 1 direct, secties 2–4 token-voor-token ─────────────────
    # Sectie 1 én de streamende secties 2–4 staan in één witte kaart
    # (.st-key-eduplan-stream-card in MAIN_CSS), zodat de tekst meteen op wit
    # verschijnt i.p.v. op de roze pagina-achtergrond.
    profiel_ph = None  # placeholder; echte st.empty()-slot wordt verderop binnen de kaart aangemaakt
    captured = {"section1": None, "final_html": None, "warning": None}
    exp = None

    def _delta_gen():
        for raw in resp.iter_lines(decode_unicode=True):
            if not raw:
                continue
            msg = json.loads(raw)
            soort = msg["type"]
            if soort == "section1":
                captured["section1"] = msg["html"]
                profiel_ph.markdown(msg["html"], unsafe_allow_html=True)
            elif soort == "warning":
                captured["warning"] = msg["html"]
            elif soort == "delta":
                yield msg["text"]
            elif soort == "final_html":
                captured["final_html"] = msg["html"]

    try:
        resp = requests.post(
            "http://localhost:8000/explain_risk_stream",
            json={
                "student": row.to_dict(),
                "prediction": result["prediction"],
                "probability": result["probability"],
                "use_default_model": gebruik_default,
                "imputed_columns": vul_log,
            },
            stream=True,
            timeout=60,
        )
        if resp.status_code != 200:
            snippet = resp.text[:300] if resp.text else "(geen response body)"
            exp = _error_html(
                f"Backend fout (HTTP {resp.status_code})",
                f"Het /explain_risk_stream endpoint gaf een fout terug. Bekijk de backend-log voor details. Response: {snippet}",
            )
        else:
            with st.container(key="eduplan-stream-card"):
                profiel_ph = st.empty()
                st.write_stream(_delta_gen())
            if captured["warning"] is not None:
                exp = (captured["section1"] or "") + captured["warning"]
            elif captured["final_html"] is not None:
                exp = (captured["section1"] or "") + captured["final_html"]
            else:
                exp = _error_html(
                    "Onverwachte response van backend",
                    "De stream eindigde zonder volledige begeleidingstekst (geen final_html).",
                )
    except requests.Timeout:
        exp = _error_html(
            "Time-out bij genereren EduPlan",
            "De backend reageerde niet binnen 60 seconden. Mogelijk is het LLM-model traag of de Anthropic API onbereikbaar.",
        )
    except requests.ConnectionError:
        exp = _error_html(
            "Backend niet bereikbaar",
            "Geen verbinding met http://localhost:8000. Controleer of de FastAPI backend draait (./1_start_fastapi.sh).",
        )
    except requests.RequestException as e:
        exp = _error_html("Netwerkfout", str(e))

    fi_dict = f_fi.result()
    pool.shutdown()
    fi_str = ", ".join(f"{k}: {v:.2f}" for k, v in fi_dict.items()) if fi_dict else "Niet beschikbaar."

    def _safe_value(col: str, cast_fn):
        if col in vul_set or col not in row.index:
            return "niet beschikbaar"
        return cast_fn(row[col])

    analyse = {
        "naam": naam,
        "opleiding": row.get("Opleiding", "—"),
        "klas": row.get("Klas", "—"),
        "mentor": row.get("Mentor", "—"),
        "studentnummer": int(row["Studentnummer"]) if "Studentnummer" in row.index else "—",
        "leeftijd": _safe_value("StudentAge", int),
        "ongeoorloofd_verzuim": _safe_value("absence_unauthorized", float),
        "geoorloofd_verzuim": _safe_value("absence_authorized", float),
        "probability": result["probability"],
        "explanation": exp,
        "feature_importance": fi_str,
        "feature_importance_dict": fi_dict,
    }
    analyse["docx"] = _build_word_doc(analyse)
    st.session_state.laatste_analyse = analyse
    st.session_state.eduplan_genereren = False


# ─────────────────────────────────────────────
# Modeltraining
# ─────────────────────────────────────────────


def _start_training() -> None:
    upload_df = st.session_state.trainings_df
    payload = {
        "data": upload_df.to_dict(orient="records"),
        "dropout_column": "Dropout",
    }
    try:
        resp = requests.post("http://localhost:8000/train_model", json=payload, timeout=10)
        resultaat = resp.json().get("status")
        if resultaat in ("started", "already_running"):
            st.session_state.training_status = "training"
            st.rerun()
    except Exception as e:
        st.error(f"Kan training niet starten: {e}")


def _reset_model() -> None:
    try:
        requests.delete("http://localhost:8000/reset_model", timeout=10)
        st.session_state.training_status = "idle"
        st.session_state.training_message = ""
        st.session_state.model_is_custom = False
        st.session_state.risicostudenten = []
        st.session_state.filter_key = None
        st.rerun()
    except Exception as e:
        st.error(f"Reset mislukt: {e}")


def _show_training_panel() -> None:
    """Toont de trainings-UI op het startscherm wanneer historische data is geüpload."""
    status = st.session_state.training_status

    if status == "idle":
        st.info(
            "Je data bevat een **Dropout-kolom** met historische uitvalgegevens. "
            "Train het model op jouw eigen data voor instelling-specifieke voorspellingen."
        )
        if st.button("Train model op jouw data", type="primary", use_container_width=True):
            _start_training()

    elif status == "training":
        # Bewaar starttijd in session state — overleeft Streamlit-reruns
        if "training_start_time" not in st.session_state:
            st.session_state.training_start_time = time.time()

        verstreken = time.time() - st.session_state.training_start_time
        minuten, seconden = divmod(int(verstreken), 60)
        tijdtekst = f"{minuten}m {seconden}s" if minuten else f"{seconden}s"

        stappen = [
            (0, "Gegevens laden en features valideren…"),
            (5, "Hyperparameterraster opbouwen (24 combinaties × 5-fold CV)…"),
            (15, "Modellen fitten — dit duurt het langst…"),
            (45, "Laatste fits afronden en beste model selecteren…"),
        ]
        fase = next(
            (msg for drempel, msg in reversed(stappen) if verstreken >= drempel),
            stappen[0][1],
        )

        with st.status("Model wordt getraind…", expanded=True):
            st.write("GridSearchCV doorzoekt hyperparameters. Dit duurt doorgaans 15–60 seconden.")
            st.markdown(f"⏱ **Verstreken tijd:** {tijdtekst}")
            st.caption(f"Fase: {fase}")

        # Eén statuscheck per render — geen blokkerende loop
        try:
            data = requests.get("http://localhost:8000/train_status", timeout=5).json()
        except Exception:
            data = {"status": "training"}

        if data["status"] == "done":
            totaal = int(time.time() - st.session_state.pop("training_start_time", time.time()))
            st.session_state.training_status = "done"
            st.session_state.training_message = f"{data['message']} (getraind in {totaal}s)"
            st.session_state.model_is_custom = True
            st.session_state.risicostudenten = []
            st.session_state.filter_key = None
            st.rerun()
        elif data["status"] == "failed":
            st.session_state.pop("training_start_time", None)
            st.session_state.training_status = "failed"
            st.session_state.training_message = data["message"]
            st.rerun()
        else:
            # Nog bezig — wacht kort en herlaad
            time.sleep(2)
            st.rerun()

    elif status == "done":
        bericht = st.session_state.get("training_message", "")
        st.success(f"Instellingsmodel actief. {bericht} Voorspellingen worden gedaan met jouw eigen getrainde model.")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Terugzetten naar standaardmodel", use_container_width=True):
                _reset_model()
        with col2:
            st.caption("Het model blijft actief na herstart van de backend.")

    elif status == "failed":
        bericht = st.session_state.get("training_message", "")
        st.error(f"Training mislukt: {bericht}")
        if st.button("Opnieuw proberen", use_container_width=True):
            st.session_state.training_status = "idle"
            st.rerun()


# ─────────────────────────────────────────────
# Startscherm
# ─────────────────────────────────────────────


def show_start_screen():
    st.markdown(START_CSS, unsafe_allow_html=True)
    st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)

    # ── Titel ──
    st.markdown(
        """
        <div style="text-align:center; margin-bottom:22px; margin-top:4px; font-family:'General Sans',sans-serif;">
            <h1 style="font-size:3.2rem; font-weight:600; line-height:1.15; margin-bottom:24px; margin-top:14px; padding:8;">
                Welkom bij de<br>Uitnodigingsregel
            </h1>
            <p style="font-size:1.3rem; font-weight:500; color:#333; margin-top:4px; padding:0;">
                op tijd de juiste lerenden uitnodigen
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    _, col_m, _ = st.columns([1, 4, 1])
    with col_m:
        # ── Beschrijvingstekst ──
        st.markdown(
            """<p style="text-align:center; font-size:1rem; font-weight:400; color:#333; margin-top:20px; margin-bottom:30px;
                         font-family:'General Sans',sans-serif; line-height:1.6;">
                Voeg je eigen dataset met lerenden toe, om te zien of er nú lerenden zijn
                die mogelijk risico lopen om uit te vallen. We brengen ze voor jou in beeld.
            </p>""",
            unsafe_allow_html=True,
        )

        # ── Upload 1: historische trainingsdata ──
        st.markdown(
            "<p style='font-size:0.9rem; font-weight:600; color:#555; margin-bottom:4px;"
            'font-family:"General Sans",sans-serif;\'>Historische data (voor modeltraining)</p>',
            unsafe_allow_html=True,
        )
        trainings_file = st.file_uploader(
            "Historische data",
            type=["csv", "xlsx"],
            label_visibility="collapsed",
            key="upload_trainingsdata",
        )
        if trainings_file is not None and trainings_file.name != st.session_state.trainings_filename:
            _verwerk_trainingsdata(trainings_file)

        # ── Modeltraining (alleen bij geüploade historische data met Dropout-kolom) ──
        if st.session_state.heeft_dropout_kolom:
            _show_training_panel()
            st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

        # ── Upload 2: huidig cohort ──
        st.markdown(
            "<p style='font-size:0.9rem; font-weight:600; color:#555; margin-bottom:4px;"
            'font-family:"General Sans",sans-serif;\'>Huidig cohort (studenten om te ranken)</p>',
            unsafe_allow_html=True,
        )
        predictie_file = st.file_uploader(
            "Huidig cohort",
            type=["csv", "xlsx"],
            label_visibility="collapsed",
            key="upload_predictiedata",
        )
        if predictie_file is not None and predictie_file.name != st.session_state.predictie_filename:
            _verwerk_predictiedata(predictie_file)

        # ── Demo-data checkbox ──
        demo_nieuw = st.checkbox(
            "Of gebruik de synthetische demo-data",
            value=st.session_state.gebruik_demo_data,
        )
        if demo_nieuw != st.session_state.gebruik_demo_data:
            st.session_state.gebruik_demo_data = demo_nieuw
            if demo_nieuw:
                st.session_state.predictie_df = None
                st.session_state.predictie_filename = ""
                st.session_state.filter_key = None
                st.session_state.risicostudenten = []
                st.session_state.laatste_analyse = None

        st.markdown("<div style='height:26px'></div>", unsafe_allow_html=True)

        # ── START-knop ──
        data_beschikbaar = st.session_state.gebruik_demo_data or st.session_state.predictie_df is not None
        if st.button(
            "START DE UITNODIGINGSREGEL",
            type="primary",
            use_container_width=True,
            disabled=not data_beschikbaar,
        ):
            st.session_state.page = "main"
            st.session_state.selected_opleiding = "Alle"
            st.session_state.selected_klas = "Alle"
            st.rerun()

        if not data_beschikbaar:
            st.caption("Upload een databestand of selecteer de synthetische demo-data om te starten.")

    st.markdown("<div style='height:20px'></div>", unsafe_allow_html=True)

    # ── Opleidingen-pills ──
    _, col_pills, _ = st.columns([0.5, 6, 0.5])
    with col_pills:
        ZICHTBAAR = 4
        eerste_rij = QUICK_OPLEIDINGEN[:ZICHTBAAR]
        rest = QUICK_OPLEIDINGEN[ZICHTBAAR:]

        pill_cols = st.columns(ZICHTBAAR + 1)
        for i, opl in enumerate(eerste_rij):
            with pill_cols[i]:
                if st.button(opl, key=f"pill_{opl}", use_container_width=True):
                    _pill_klik(opl)
                    st.rerun()
        with pill_cols[ZICHTBAAR]:
            if not st.session_state.toon_alle_opleidingen:
                if st.button("Meer ↓", key="pill_meer", use_container_width=True):
                    st.session_state.toon_alle_opleidingen = True
                    st.rerun()
            else:
                if st.button("Minder ↑", key="pill_minder", use_container_width=True):
                    st.session_state.toon_alle_opleidingen = False
                    st.rerun()

        if st.session_state.toon_alle_opleidingen and rest:
            extra_cols = st.columns(min(len(rest), 5))
            for i, opl in enumerate(rest):
                with extra_cols[i % len(extra_cols)]:
                    if st.button(opl, key=f"pill_extra_{opl}", use_container_width=True):
                        _pill_klik(opl)
                        st.rerun()


# ─────────────────────────────────────────────
# Hoofdscherm — header
# ─────────────────────────────────────────────


def _render_header():
    tab = st.session_state.actieve_tab

    col_ceda, col_terug, col_ur, col_ep = st.columns([3, 1.2, 2.2, 1.2])

    with col_ceda:
        st.markdown(
            "<img src='app/static/npuls-logo.svg' alt='Npuls logo'"
            " style='height:48px;width:auto;display:block;padding:6px 0;'>",
            unsafe_allow_html=True,
        )
    with col_terug:
        st.markdown("<div class='nav-terug'>", unsafe_allow_html=True)
        if st.button("← TERUG", key="nav_terug", use_container_width=True):
            st.session_state.page = "start"
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)
    with col_ur:
        klasse_ur = "nav-actief" if tab == "uitnodigingsregel" else "nav-inactief"
        st.markdown(f"<div class='{klasse_ur}'>", unsafe_allow_html=True)
        if st.button("UITNODIGINGSREGEL", key="nav_ur", use_container_width=True):
            st.session_state.actieve_tab = "uitnodigingsregel"
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)
    with col_ep:
        klasse_ep = "nav-actief" if tab == "eduplan" else "nav-inactief"
        st.markdown(f"<div class='{klasse_ep}'>", unsafe_allow_html=True)
        if st.button("EDUPLAN", key="nav_ep", use_container_width=True):
            st.session_state.actieve_tab = "eduplan"
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)


# ─────────────────────────────────────────────
# Hoofdscherm — kaart-header (opleiding + klas + potlood)
# ─────────────────────────────────────────────


def _render_card_header():
    opl = st.session_state.selected_opleiding

    if st.session_state.toon_zoekbalk:
        st.markdown("<div class='card-zoek'>", unsafe_allow_html=True)
        col_s, col_b = st.columns([5, 1])
        with col_s:
            zoek = st.text_input(
                "zoek",
                placeholder="🔍  Zoek een andere opleiding (b.v. Economie, Techniek)...",
                label_visibility="collapsed",
                key="card_zoek_input",
            )
        with col_b:
            if st.button("ZOEK", type="primary", use_container_width=True, key="card_zoek_btn"):
                if zoek.strip():
                    st.session_state.selected_opleiding = _zoek_opleiding(zoek.strip())
                st.session_state.selected_klas = "Alle"
                st.session_state.toon_zoekbalk = False
                st.session_state.filter_key = None
                st.session_state.risicostudenten = []
                st.session_state.laatste_analyse = None
                st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    else:
        klassen = _klassen_voor(opl)

        if st.session_state.selected_klas not in klassen:
            st.session_state.selected_klas = "Alle"

        col_t, col_k, col_sp, col_p = st.columns([3, 2.5, 2.5, 0.5])

        with col_t:
            badge = (
                " <span style='font-size:0.7rem; background:#e8f5e9; color:#2e7d32;"
                " border-radius:4px; padding:2px 6px; vertical-align:middle;"
                " font-weight:500;'>instellingsmodel</span>"
                if st.session_state.get("model_is_custom")
                else ""
            )
            st.markdown(
                f"<h3 style='font-weight:500; margin:0; padding:6px 0;"
                f'font-family:"General Sans",sans-serif;\'>{opl}{badge}</h3>',
                unsafe_allow_html=True,
            )

        with col_k:
            klas_idx = klassen.index(st.session_state.selected_klas) if st.session_state.selected_klas in klassen else 0
            gekozen_klas = st.selectbox(
                "klas",
                klassen,
                index=klas_idx,
                label_visibility="collapsed",
                format_func=lambda x: f"KLAS:  {x}" if x != "Alle" else "KLAS:  Alle",
                key="klas_dropdown",
            )
            if gekozen_klas != st.session_state.selected_klas:
                st.session_state.selected_klas = gekozen_klas
                st.session_state.filter_key = None
                st.session_state.laatste_analyse = None
                st.rerun()

        with col_p:
            st.markdown("<div class='potlood-btn'>", unsafe_allow_html=True)
            if st.button("🖊️", key="potlood"):
                st.session_state.toon_zoekbalk = True
                st.rerun()
            st.markdown("</div>", unsafe_allow_html=True)


# ─────────────────────────────────────────────
# Hoofdscherm — banner  "Toon mij X lerenden…"
# ─────────────────────────────────────────────


def _render_banner():
    risico = st.session_state.risicostudenten

    label = f"<u><b>{st.session_state.top_n}</b></u>" if risico else "···"

    st.markdown(
        f"""<div style="background:{TERRACOTTA}; border-radius:12px; padding:14px 24px;
                        color:white; font-size:16px; font-weight:500;
                        text-align:center; margin:16px 0;
                        font-family:'General Sans',sans-serif;">
            Toon mij {label} lerenden met het hoogste risico om uit te vallen.
        </div>""",
        unsafe_allow_html=True,
    )

    if risico:
        max_n = len(risico)
        nieuw_n = st.slider(
            "Aantal te tonen",
            min_value=1,
            max_value=max_n,
            value=min(st.session_state.top_n, max_n),
            label_visibility="collapsed",
            key="top_n_slider",
        )
        if nieuw_n != st.session_state.top_n:
            st.session_state.top_n = nieuw_n
            st.rerun()


# ─────────────────────────────────────────────
# Hoofdscherm — barchart
# ─────────────────────────────────────────────


def _render_barchart():
    risico = st.session_state.risicostudenten

    if not risico:
        st.markdown(
            f"""<div style="background:{ROZE_LICHT}; border-radius:14px;
                            height:260px; display:flex; align-items:center;
                            justify-content:center; color:#bbb; font-size:1.4rem;
                            margin:8px 0; letter-spacing:0.1em;">
                · · ·
            </div>""",
            unsafe_allow_html=True,
        )
        return

    top_n = st.session_state.top_n
    top = risico[:top_n]
    n = len(top)

    namen = [row["Naam"] for row, _ in reversed(top)]
    kansen = [result["probability"] for _, result in reversed(top)]

    def terracotta(i, total):
        t = i / max(total - 1, 1)
        r = int(0xA0 + (0xDF - 0xA0) * t)
        g = int(0x55 + (0x9A - 0x55) * t)
        b = int(0x35 + (0x75 - 0x35) * t)
        return f"rgb({r},{g},{b})"

    kleuren = [terracotta(i, n) for i in range(n)]

    fig = go.Figure(
        go.Bar(
            x=kansen,
            y=namen,
            orientation="h",
            marker_color=kleuren,
            text=[f"{k:.0%}" for k in kansen],
            textposition="outside",
            textfont=dict(size=14, color="#aaa"),
        )
    )
    x_max = max(kansen) * 1.35 if kansen else 1.0
    fig.update_layout(
        xaxis=dict(range=[0, x_max], showgrid=False, showticklabels=False, zeroline=False),
        yaxis=dict(showgrid=False, tickfont=dict(size=14, color="#1a1a1a")),
        plot_bgcolor="white",
        paper_bgcolor="white",
        margin=dict(l=20, r=60, t=12, b=12),
        height=max(180, 60 + 56 * n),
        showlegend=False,
    )
    st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})


# ─────────────────────────────────────────────
# Hoofdscherm — EduPlan-sectie
# ─────────────────────────────────────────────


def _render_eduplan_sectie():
    risico = st.session_state.risicostudenten
    top_n = st.session_state.top_n

    if not risico:
        st.info("Laad eerst studenten via het UITNODIGINGSREGEL-tabblad.")
        return

    top = risico[:top_n]

    st.markdown(
        "<p style='font-size:14px; color:#000; margin:8px 0 4px 0;"
        'font-family:"General Sans",sans-serif;\'>'
        "<b>Selecteer een lerende voor de uitleg van diens uitvalrisico</b></p>",
        unsafe_allow_html=True,
    )

    opties = [f"{row['Naam'].upper()} — UITVALKANS: {result['probability']:.0%}" for row, result in top]

    col_sel, col_btn = st.columns([4, 2])
    with col_sel:
        st.markdown("<div class='student-sel'>", unsafe_allow_html=True)
        idx = st.selectbox(
            "student",
            options=range(len(opties)),
            format_func=lambda i: opties[i],
            label_visibility="collapsed",
            key="student_selector",
        )
        st.markdown("</div>", unsafe_allow_html=True)

    with col_btn:
        if st.button("TOON EDUPLAN", type="primary", use_container_width=True):
            st.session_state.geselecteerde_student = idx
            st.session_state.eduplan_genereren = True
            st.session_state.laatste_analyse = None
            st.session_state.actieve_tab = "eduplan"
            st.rerun()

    if st.session_state.eduplan_genereren:
        naam = html.escape(top[st.session_state.geselecteerde_student][0]["Naam"])
        st.markdown(
            f"""<div style="display:flex; align-items:center; margin-top:16px;
                            font-family:'General Sans',sans-serif;">
                <span style="font-size:28px; font-weight:700;">🚦</span>&nbsp;&nbsp;
                <span style="font-size:2.0rem; font-weight:500;">EduPlan | {naam}</span>
            </div>""",
            unsafe_allow_html=True,
        )
        # Sectie 1 verschijnt direct, secties 2–4 streamen live binnen deze functie.
        _genereer_eduplan()
        st.rerun()

    if st.session_state.laatste_analyse:
        _render_eduplan_content()


def _render_eduplan_content():
    analyse = st.session_state.laatste_analyse
    naam = html.escape(analyse["naam"])

    with st.container(border=True):
        st.markdown(
            f"""<div style="display:flex; background:white; border-radius:16px; align-items:center; 
                            font-family:'General Sans',sans-serif;">
                <div style="background:white; color:#1a1a1a; border-radius:16px;
                            width:68px; height:59px; display:flex; align-items:center;
                            justify-content:center; font-size:28px; font-weight:700;
                            flex-shrink:1;">🚦</div>
                <span style="background:white; font-size:2.0rem; font-weight:500;">EduPlan | {naam}</span>
            </div>""",
            unsafe_allow_html=True,
        )

    st.markdown("<div style='height:26px'></div>", unsafe_allow_html=True)

    st.markdown(
        f'<div style=\'font-family:"General Sans",sans-serif; font-size:15px; '
        f"line-height:1.85; background:white; border-radius:16px; "
        f"padding:28px 32px;'>"
        f"{analyse['explanation']}"
        f"</div>",
        unsafe_allow_html=True,
    )

    st.markdown("<div style='height:16px'></div>", unsafe_allow_html=True)

    _, col_p, col_d = st.columns([4, 1.5, 1.5])

    with col_d:
        st.download_button(
            label="DOWNLOAD",
            data=analyse["docx"],
            file_name=f"EduPlan_{naam.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="download_btn",
        )


# ─────────────────────────────────────────────
# Hoofdscherm — footer
# ─────────────────────────────────────────────


def _render_footer():
    st.markdown("<div style='height:32px'></div>", unsafe_allow_html=True)
    st.markdown(
        """<hr style="border:none; border-top:2px solid rgba(0,0,0,0.32); margin:0 0 20px 0;">
        <p style="text-align:center; font-size:0.75rem; font-weight:500; color:gray; font-family:'General Sans',sans-serif;">
            &#169; &#9432; 2026 Op deze analytics tool is de Creative Commons ShareAlike
            Naamsvermelding 4.0-licentie van toepassing. Maak bij gebruik van dit werk
            vermelding van de volgende referentie: AI en data waarde(n)vol inzetten: CEDA.
            Uitnodigingsregel – EduPlan. Utrecht: Npuls
        </p>""",
        unsafe_allow_html=True,
    )


# ─────────────────────────────────────────────
# Hoofdscherm — samenstellen
# ─────────────────────────────────────────────


def show_main_screen():
    st.markdown(MAIN_CSS, unsafe_allow_html=True)

    _render_header()

    dff = _gefilterde_df()

    if not dff.empty:
        _run_voorspelling(dff)

    with st.container(border=True):
        _render_card_header()

        tab = st.session_state.actieve_tab

        _render_banner()

        if tab == "uitnodigingsregel":
            _render_barchart()
        else:
            _render_eduplan_sectie()

    # _render_footer()


# ─────────────────────────────────────────────
# Router
# ─────────────────────────────────────────────

if st.session_state.page == "start":
    show_start_screen()
else:
    show_main_screen()

with bottom():
    _col_txt, _col_btn = st.columns([5, 1])
    with _col_txt:
        st.markdown(
            """<p style="text-align:center; font-size:0.5rem; font-weight:500; color:#1a1a1a;
                    font-family:'General Sans',sans-serif; margin:1px 0;"><br>
                    <img src="https://mirrors.creativecommons.org/presskit/icons/cc.svg" alt="" style="max-width: 2em;max-height:3em;margin-left: .2em;"><img src="https://mirrors.creativecommons.org/presskit/icons/by.svg" alt="" style="max-width: 2em;max-height:3em;margin-left: .2em;"> Op deze analytics tool is de Creative Commons ShareAlike
            Naamsvermelding 4.0-licentie van toepassing. <br>Maak bij gebruik van dit werk
            vermelding van de volgende referentie: AI en data waarde(n)vol inzetten: CEDA
            2026 Uitnodigingsregel – EduPlan. Utrecht: Npuls
        </p>""",
            unsafe_allow_html=True,
        )
    with _col_btn:
        if st.session_state.page != "start":
            try:
                with open("docs/model_analysis.html", "rb") as _f:
                    st.download_button(
                        label="📊 Modelevaluatie",
                        data=_f.read(),
                        file_name="model_analysis.html",
                        mime="text/html",
                        use_container_width=True,
                        help="Download het modelevaluatie-rapport (HTML)",
                    )
            except FileNotFoundError:
                pass
